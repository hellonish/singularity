"""
Loader for DOCX files — extracts paragraph and table text.

Requires: pip install python-docx
"""
import asyncio
from typing import List

from models.models import Document
from .base import BaseLoader


class DocxLoader(BaseLoader):
    """
    Loader for .docx (Microsoft Word) files using python-docx.
    """

    async def load(self, source: str, **kwargs) -> List[Document]:
        """
        Load text content from a .docx file.

        Args:
            source: Absolute path to a .docx file.
        """
        try:
            def parse_docx():
                import docx  # python-docx

                doc = docx.Document(source)
                parts: list[str] = []

                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            parts.append(" | ".join(cells))

                return "\n\n".join(parts)

            content = await asyncio.to_thread(parse_docx)

            return [
                Document(
                    id=source,
                    content=content,
                    metadata={"source": "docx", "path": source},
                )
            ]
        except Exception as e:
            print(f"Error in DocxLoader: {e}")
            return []
